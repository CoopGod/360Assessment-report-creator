import csv
import sys
import pandas as pd
import matplotlib.pyplot as plt
from PIL import Image
import docx
from docx.shared import Inches


class Review:
    def __init__(self, focusName, focusRelation, reviewName, skillList, importanceList):
        self.focusName = focusName
        self.focusRelation = focusRelation
        self.reviewName = reviewName
        self.skillList = skillList,
        self.importanceList = importanceList


def main():
    # Ensure proper CMD Line Arg
    if len(sys.argv) > 3:
        print("Error!")
        return 1
    revieweesList = reviewees()
    for review in revieweesList:
        assessments = dataParse(review)
        df = createDataFrame(assessments)
        plotData(df, review)


# function to find all people of focus
def reviewees():
    schedule = sys.argv[1]
    with open(schedule, "r", encoding="utf8") as schedule:
        reader = csv.reader(schedule)
        reviewees = []
        for row in reader:
            if row[3] not in reviewees:
                reviewees.append(row[3])
    # remove header
    reviewees.pop(0)
    return reviewees


# function to create Review object regarding person of focus
def dataParse(reviewee):
    # Upload CSV files
    schedule = sys.argv[1]
    # open CSV file and run through data. adding to nested list
    assessments = []
    with open(schedule, "r", encoding="utf8") as schedule:
        reader = csv.reader(schedule)
        # Move through rows including header
        rowCount = 0
        employeeCount = 0
        for row in reader:
            # if person of focus
            if row[3] == reviewee or rowCount == 0:
                # if not header move through and assign based on documentation
                if rowCount == 0:
                    headerList = []
                    headerList = row
                    rowCount += 1
                    continue
                # priliminary setup
                name = row[1]
                focus = row[3]
                # change relation name to described documentation
                if row[4] == "Peer" or row[4] == "Employee":
                    employeeCount += 1
                    relation = f"Employee {employeeCount}"
                else:
                    relation = row[4]
                skills = []
                importance = []
                # if column is skill (based on number beginning the question) assign it to its list. If question starts
                # with opera... then assign it to importance
                itemCount = 0
                previousHeader = None
                for item in row:
                    currentHeader = headerList[itemCount]
                    if currentHeader[:1].isdigit():
                        skills.append(f"{currentHeader[:2]}:{item}")
                        previousHeader = currentHeader
                    elif currentHeader[:2] == "Op":
                        importance.append(f"{previousHeader[:2]}:{item}")
                    itemCount += 1
                # create Review object for later use
                assessments.append(
                    Review(focus, relation, name, skills, importance))
                rowCount += 1            
    return assessments


# function to create data frame regarding person of focus
def createDataFrame(assessments):
    # init all columns in data frame
    columnList = []
    for i in range(1, 38):
        if i < 10:
            number = f"{i}."
        else:
            number = str(i)
        columnList.append(number)
    df = pd.DataFrame(columns=columnList)
    
    # Log all Skill Ratings!
    skillReviewNum = 0
    axisRelations = {}
    for review in assessments:
        skillValues = {}
        for skills in review.skillList:
            for skill in skills:
                # move through each skill and append it to the corresponding data frame column
                data = int(skill[3:])
                header = str(skill[:2])
                skillValues[header] = data
        # add all review data to df
        df = df.append(skillValues, ignore_index=True)
        # change axis names to reviewers relationship and indicate which they are rating (skill vs importance)
        axisRelations[skillReviewNum] = review.focusRelation + "-SKL"
        skillReviewNum += 1

    # log all Importance Ratings!
    impReviewNum = 0
    for review in assessments:
        impValues = {}
        for imps in review.importanceList:
            # move through each importance rating and append it to the corresponding data frame column
            data = int(imps[3:])
            header = str(imps[:2])
            impValues[header] = data
        # add all review data to df
        df = df.append(impValues, ignore_index=True)
        # change axis names to reviewers relationship and indicate which they are rating (skill vs importance)
        axisRelations[impReviewNum +
                      skillReviewNum] = review.focusRelation + "-IMP"
        impReviewNum += 1
    # rename and sort alphebetically
    df = df.rename(index=axisRelations)
    df = df.sort_index(ascending=False)

    # create legend for all graphs
    try:
        document = docx.Document('legend.docx')
    except:
        document = docx.Document()
    for review in assessments:
        document.add_heading(f'{review.focusName}', 0)
        document.add_paragraph(f"{review.reviewName}: {review.focusRelation}")
    document.save('legend.docx')

    return df


# function to upload/display graphs using data frame
def plotData(df, reviewee):
    # get different colors
    colorDictionary = {0: '#E62421', 2: '#344EE3', 4: '#34B8E3', 6: '#A1FF37', 8: '#14CFFF', 10: '#3634E3', 12: '#DF46FA', 14: '#E34234', 16: '#FB9F3A',
                       18: '#FAF6D1', 20: '#E6AF8C', 22: '#9AC3E6', 24: '#C4B5FA', 26: '#FADA82', 28: '#E6AFAC', 30: '#9EE6DA', 32: '#60E651', 34: '#FA72F7'}  # TODO
    cmap = []
    bars = len(df.index)
    for num in range(bars):
        if num % 2 != 0:
            cmap.append(cmap[num - 1])
        else:
            cmap.append(colorDictionary[num])
    # save file in report
    # document setup
    document = docx.Document()
    document.add_heading(f'360 Assessment - {reviewee}', 0)
    # find complete questions in header of original CSV
    schedule = sys.argv[1]
    with open(schedule, "r", encoding="utf8") as schedule:
        reader = csv.DictReader(schedule)
        fieldnames = reader.fieldnames
        headers = []
        for field in fieldnames:
            field = field.replace("Skilled:", "")
            headers.append(field)
    # create plot and add image of it to docx
    for i in range(1, len(df.columns) + 1):
        iteration = str(i) + "."  # to skirt the "." issue with 1-9
        ax = df.plot.bar(y=iteration[:2], color=cmap)  # fuck ya
        # pretty graphs <3
        ax.set_ylim(ymax=5.25)
        ax.get_legend().remove()
        plt.yticks(rotation=90)
        plt.yticks([0, 1, 2, 3, 4, 5], ['I', 'D', 'BA', 'C', 'VG', 'E'])
        # create table to imitate textbox + add question info/text to document
        table = document.add_table(rows=1, cols=2)
        row = table.rows[0].cells
        row[1].text = headers[3+(i*2)]
        # add image
        plt.savefig('temp.jpg', bbox_inches='tight')
        imageFile = Image.open('temp.jpg')
        # add whitespace to image so rotating isn't wonky
        background = Image.new("RGB", (550, 550), (255, 255, 255))
        background.paste(imageFile)
        background = background.rotate(270)
        background.save('temp.jpg')
        run = row[0].paragraphs[0].add_run()
        run.add_picture('temp.jpg', width=Inches(4))
        # force page break every two entries
        if i % 2 == 0:
            document.add_page_break()
        # save that sweet, sweet memory
        imageFile.close()
        plt.close()
    # save file under person of focus' name
    document.save(f'{reviewee}-report.docx')


if __name__ == "__main__":
    main()
