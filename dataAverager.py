# import packages
import csv
import sys
import pandas as pd
import matplotlib.pyplot as plt
from PIL import Image
import docx
from docx.shared import Inches
# import function from helper file
from parse import *

class Review:
    def __init__(self, focusRelation, skillList, importanceList):
        self.focusRelation = focusRelation
        self.skillList = skillList,
        self.importanceList = importanceList    


def main():
    # Ensure proper CMD Line Arg
    if len(sys.argv) > 2:
        print("Error!")
        return 1
    assessments = dataParse()
    df = createDataFrame(assessments)
    plotData(df)


# function to create Review object to then be input into dataframe
def dataParse():
    # Upload CSV files
    scheduleArg = sys.argv[1]
    # open CSV file and run through data. adding to nested list
    assessments = []
    with open(scheduleArg, "r", encoding="utf8") as schedule:
        reader = csv.reader(schedule)
        # Move through rows including header        
        rowCount = employeeCount = selfCount = boardCount = peerCount = 0
        headerList = []
        employeeRows = []
        selfRows = []
        boardRows = []
        peerRows = []
        for row in reader:
            # if not header move through and assign based on documentation
            if rowCount == 0:
                headerList = row
                rowCount += 1
                continue # check this to make sure it works
            # priliminary setup
            name = row[1]
            focus = row[3]

            # If employee, bring to different loop and average those scores
            if row[4] == "Employee":
                employeeCount += 1
                employeeRows.append(rowCount)
            # else add the rest to their own columns
            elif row[4] == "Self":
                selfCount += 1
                selfRows.append(rowCount)
            elif row[4] == "Peer":
                peerCount += 1
                peerRows.append(rowCount)
            else:
                boardCount += 1
                boardRows.append(rowCount)
            rowCount += 1
                
    schedule.close()

    # parse through all employees and add their stats
    infoArray = parse(scheduleArg, employeeRows, headerList, employeeCount, 'Employee')
    assessments.append(Review(infoArray[0], infoArray[1], infoArray[2]))

    # parse through all Self reviews and add their stats
    infoArray = parse(scheduleArg, selfRows, headerList, selfCount, 'Self')
    assessments.append(Review(infoArray[0], infoArray[1], infoArray[2]))
    print(infoArray[0])

    # parse through all Peer reviews and add their stats
    infoArray = parse(scheduleArg, peerRows, headerList, peerCount, 'Peer')
    assessments.append(Review(infoArray[0], infoArray[1], infoArray[2]))
            
    # parse through all boards and add their stats
    infoArray = parse(scheduleArg, boardRows, headerList, boardCount, 'Board')
    assessments.append(Review(infoArray[0], infoArray[1], infoArray[2]))

    return assessments


# function to create data frame regarding person of focus
def createDataFrame(assessments):
    # init all columns in data frame
    columnList = []
    for i in range(1, 38):
        if i < 10:
            number = f"{i}."
        else:
            number = str(i)
        columnList.append(number)
    df = pd.DataFrame(columns=columnList)

    # Log all Skill Ratings!
    skillReviewNum = 0
    axisRelations = {}
    for review in assessments:
        skillValues = review.skillList
        # add all review data to df
        temp = pd.DataFrame(skillValues, columns=columnList)
        df = df._append(temp, ignore_index=True)
        # change axis names to reviewers relationship and indicate which they are rating (skill vs importance)
        axisRelations[skillReviewNum] = review.focusRelation + "-SKL"
        skillReviewNum += 1

    # log all Importance Ratings!
    impReviewNum = 0
    for review in assessments:
        impValues = review.importanceList
        # add all review data to df
        temp = pd.DataFrame([impValues], columns=columnList)
        df = df._append(temp, ignore_index=True)
        # change axis names to reviewers relationship and indicate which they are rating (skill vs importance)
        axisRelations[impReviewNum + skillReviewNum] = review.focusRelation + "-IMP"
        impReviewNum += 1

    # rename and sort alphebetically
    df = df.rename(index=axisRelations)
    df = df.sort_index(ascending=False)
    df.plot.bar(y="1.")
    # add averages to separate document
    df.to_csv('legend.csv', header=True)
    return df

# function to plot dataframe data to report
def plotData(df):
    # get different colors
    colorDictionary = {0: '#fc0403',2: '#008000',4: '#0000FF',6: '#FF7F50',8: '#FFFACD',10: '#FF69B4', 12: '#800080'}
    cmap = []
    bars = len(df.index)
    for num in range(bars):
        if num % 2 != 0:
            cmap.append(cmap[num - 1])
        else:
            cmap.append(colorDictionary[num])
    # save file in report
    # document setup
    document = docx.Document()
    document.add_heading(f'360 Assessment - Averages',0)
    # find complete questions in header of original CSV
    schedule = sys.argv[1]
    with open(schedule, "r", encoding="utf8") as schedule:
        reader = csv.DictReader(schedule)
        fieldnames = reader.fieldnames
        headers = []
        for field in fieldnames:
            field = field.replace("Skilled:", "")
            headers.append(field)
    # create plot and add image of it to docx
    for i in range(1, len(df.columns) + 1):
        iteration = str(i) + "." # to skirt the "." issue with 1-9
        ax = df.plot.bar(y=iteration[:2], color=cmap) # fuck ya
        # pretty graphs <3
        ax.set_ylim(ymax=5.25)
        ax.get_legend().remove()
        plt.yticks(rotation= 90)
        plt.yticks([0,1,2,3,4,5], ['I', 'D', 'BA', 'C', 'VG', 'E'])
        # add question info/text to document
        paragraph = document.add_paragraph(headers[3+(i*2)])
        # add image
        plt.savefig('temp.jpg', bbox_inches='tight')
        imageFile = Image.open('temp.jpg')
        # create table to imitate textbox
        table = document.add_table(rows=1, cols=2)
        row = table.rows[0].cells
        row[1].text = "SKL:  IMP:"
        # add whitespace to image so rotating isn't wonky
        background = Image.new("RGB", (550, 550), (255,255,255))
        background.paste(imageFile)
        background = background.rotate(270)
        background.save('temp.jpg')
        run = row[0].paragraphs[0].add_run()
        run.add_picture('temp.jpg', width=Inches(3))
        # force page break every two entries
        if i % 2 == 0:
            document.add_page_break()
        # save that sweet, sweet memory
        imageFile.close()
        plt.close()
    # save file under person of focus' name
    document.save('averages-report.docx')


if __name__ == "__main__":
    main()